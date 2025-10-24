import os
import re
import base64
import requests
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from markdown import markdown
from bs4 import BeautifulSoup


class DocxBuilder:
    """
    🔹 تحويل نص Mistral OCR (Markdown) إلى ملف Word مع تنسيقات كاملة.
    🔹 يدعم RTL/LTR تلقائيًا حسب اللغة (عربية أو إنجليزية).
    """

    def __init__(self, output_dir: str = "outputs"):
        self.font_name = "Arial"
        self.font_size = 12
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    # ------------------------------------------------------------
    def markdown_to_docx(self, markdown_text: str) -> str:
        """تحويل Markdown إلى Word مع دعم RTL/LTR التلقائي."""
        html = markdown(markdown_text, extensions=["tables", "fenced_code"])
        soup = BeautifulSoup(html, "html.parser")

        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        for element in soup.children:
            self._process_element(element, doc)

        # ضبط الاتجاه والخط لجميع الفقرات
        for paragraph in doc.paragraphs:
            self._apply_direction(paragraph)
            for run in paragraph.runs:
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)

        filename = f"ocr_output_{os.urandom(4).hex()}.docx"
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        return output_path

    # ------------------------------------------------------------
    def _is_arabic(self, text: str) -> bool:
        """كشف إذا كان النص عربيًا بناءً على الحروف."""
        arabic_chars = re.findall(r"[\u0600-\u06FF]", text)
        return len(arabic_chars) > len(text) * 0.3  # إذا أكثر من 30% حروف عربية

    # ------------------------------------------------------------
    def _apply_direction(self, paragraph):
        """ضبط اتجاه النص تلقائيًا (RTL أو LTR)."""
        text = paragraph.text.strip()
        if not text:
            return

        is_arabic = self._is_arabic(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_arabic else WD_ALIGN_PARAGRAPH.LEFT

        p_pr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        if is_arabic:
            p_pr.append(bidi)  # تفعيل RTL
        # إذا النص إنجليزي → لا نضيف bidi فيبقى LTR طبيعيًا

    # ------------------------------------------------------------
    def _process_element(self, element, doc):
        """تحليل وتحويل العناصر HTML إلى عناصر Word."""
        if element.name is None:
            text = element.strip()
            if text:
                p = doc.add_paragraph(text)
                self._apply_direction(p)
            return

        # العناوين
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            p = doc.add_paragraph(element.get_text(), style=f"Heading {min(level, 3)}")
            self._apply_direction(p)
            return

        # الفقرات
        if element.name == "p":
            p = doc.add_paragraph()
            self._add_inline_formatting(p, element)
            self._apply_direction(p)
            return

        # القوائم
        if element.name == "ul":
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(li.get_text(), style="List Bullet")
                self._apply_direction(p)
            return

        if element.name == "ol":
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(li.get_text(), style="List Number")
                self._apply_direction(p)
            return

        # الجداول
        if element.name == "table":
            rows = element.find_all("tr")
            if not rows:
                return
            cols = len(rows[0].find_all(["td", "th"]))
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for j, cell in enumerate(cells):
                    table.cell(i, j).text = cell.get_text(strip=True)
            return

        # الصور
        if element.name == "img":
            img_src = element.get("src")
            if not img_src:
                return
            try:
                img_data = None
                if img_src.startswith("data:image"):
                    header, encoded = img_src.split(",", 1)
                    img_data = base64.b64decode(encoded)
                else:
                    response = requests.get(img_src)
                    img_data = response.content
                img_stream = BytesIO(img_data)
                doc.add_picture(img_stream, width=Inches(5))
            except Exception:
                doc.add_paragraph("[🖼️ لم يتم تحميل الصورة]")
            return

        # فاصل أفقي
        if element.name == "hr":
            p = doc.add_paragraph("----------------------------")
            self._apply_direction(p)
            return

        # عناصر متداخلة
        for child in element.children:
            self._process_element(child, doc)

    # ------------------------------------------------------------
    def _add_inline_formatting(self, paragraph, element):
        """تحليل العناصر الداخلية (غامق/مائل/روابط...)."""
        for child in element.children:
            if isinstance(child, str):
                paragraph.add_run(child)
            elif child.name in ["strong", "b"]:
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name in ["em", "i"]:
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == "u":
                run = paragraph.add_run(child.get_text())
                run.underline = True
            elif child.name == "a":
                href = child.get("href", "#")
                run = paragraph.add_run(child.get_text())
                paragraph.add_run(f" ({href})")
            else:
                self._add_inline_formatting(paragraph, child)
